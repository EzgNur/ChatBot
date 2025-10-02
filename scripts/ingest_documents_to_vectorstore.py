"""
Belgeler klasöründeki DOCX dosyalarını vectorstore'a ekleme scripti
Alternativkraft blog makaleleriyle aynı formatta işler
"""

import sys
import os
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict
import argparse

# Proje kökünü PYTHONPATH'e ekle
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.append(PROJECT_ROOT)

try:
    from docx import Document  # python-docx
except ImportError:
    print("❌ python-docx paketi gerekli. Kurulum: pip install python-docx")
    sys.exit(1)

from vectorstore.build_store import OptimizedVectorStoreBuilder


def docx_to_text(path: Path) -> str:
    """DOCX dosyasını metin formatına çevirir"""
    try:
        doc = Document(str(path))
        parts = []
        
        # Paragrafları işle
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if not txt:
                continue
                
            style = (p.style.name or "").lower()
            if style.startswith("heading"):
                # Başlıkları markdown formatında
                level = 2
                try:
                    level = int(''.join(ch for ch in p.style.name if ch.isdigit()) or '2')
                except Exception:
                    level = 2
                parts.append(("#" * max(2, min(6, level))) + f" {txt}")
            elif style.startswith("list") or txt.startswith(("•", "-", "*")):
                # Liste öğelerini normalize et
                clean = txt.lstrip("•*- ").strip()
                parts.append(f"- {clean}")
            else:
                parts.append(txt)
        
        # Tabloları işle
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        
        return "\n".join(parts).strip()
        
    except Exception as e:
        print(f"❌ DOCX okuma hatası ({path}): {e}")
        return ""


def create_document_data(file_path: Path, title: str, url: str, content: str) -> Dict:
    """Belge verisini blog makaleleriyle aynı formatta oluşturur"""
    return {
        "id": None,  # Vectorstore builder tarafından atanacak
        "title": title,
        "content": content,
        "author": "Alternativkraft GmbH",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "url": url,
        "word_count": len(content.split()) if content else 0,
        "scraped_at": datetime.now().isoformat(),
        "scraper_type": "document_ingest",
        "content_type": "service_protocol",
        "language": "tr",
        "domain": "law_immigration_politics"
    }


def process_documents(documents_dir: str, base_url: str = "https://alternativkraft.com/") -> List[Dict]:
    """Belgeler klasöründeki DOCX dosyalarını işler"""
    
    # URL eşleştirmeleri
    url_mappings = {
        "EK 5 MULTECILIK HIZMET TANIM PROTOKOLÜ.docx": f"{base_url}service/muelteci-hukuku/",
        "Ek 5 81A Profesyonel Sürücüler için Almanya Vize Başvuru Süreci.docx": f"{base_url}service/profesyonel-sueruecue-olmak/",
        "Ek 5 81a Ön Onaylı Çalışma Vizesi Hizmet Protokolü.docx": f"{base_url}service/81a-almanya-oen-onay-vizesi/"
    }
    
    documents_dir = Path(documents_dir)
    if not documents_dir.exists():
        print(f"❌ Belgeler klasörü bulunamadı: {documents_dir}")
        return []
    
    docx_files = list(documents_dir.glob("*.docx"))
    if not docx_files:
        print("❌ Hiç DOCX dosyası bulunamadı")
        return []
    
    print(f"📄 {len(docx_files)} DOCX dosyası bulundu")
    
    processed_documents = []
    
    for file_path in docx_files:
        print(f"📖 İşleniyor: {file_path.name}")
        
        # İçeriği çıkar
        content = docx_to_text(file_path)
        if not content or len(content.strip()) < 100:
            print(f"⚠️  İçerik çok kısa veya boş: {file_path.name}")
            continue
        
        # Başlık oluştur
        title = file_path.stem.replace("_", " ").replace("-", " ").strip()
        title = " ".join(word.capitalize() for word in title.split())
        
        # URL belirle
        url = url_mappings.get(file_path.name, f"{base_url}service/")
        
        # Belge verisini oluştur
        doc_data = create_document_data(file_path, title, url, content)
        processed_documents.append(doc_data)
        
        print(f"✅ İşlendi: {title} ({doc_data['word_count']} kelime)")
    
    return processed_documents


def save_documents_json(documents: List[Dict], output_path: str = None) -> str:
    """İşlenmiş belgeleri JSON dosyasına kaydeder"""
    
    if not output_path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"data/raw/documents_{timestamp}.json"
    
    # Klasörü oluştur
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(documents, f, ensure_ascii=False, indent=2)
    
    print(f"💾 Belgeler kaydedildi: {output_path}")
    return output_path


def add_to_vectorstore(json_path: str, rebuild: bool = False):
    """Belgeleri vectorstore'a ekler"""
    
    builder = OptimizedVectorStoreBuilder()
    
    if rebuild:
        print("🚀 Rebuild modu: vectorstore baştan oluşturuluyor...")
        builder.process_clean_json_to_vectorstore(json_path)
    else:
        print("➕ Incremental modu: mevcut vectorstore'a ekleniyor...")
        
        # Mevcut vectorstore'u yükle
        vs_path = os.path.join(PROJECT_ROOT, "data", "vectorstore")
        vectorstore = builder.load_vectorstore(vs_path)
        
        if vectorstore is None:
            print("⚠️ Mevcut FAISS bulunamadı. İlk kez oluşturulacak.")
            vectorstore = builder.process_clean_json_to_vectorstore(json_path)
            return
        
        # Yeni veriyi yükle ve ekle
        documents = builder.load_clean_json_data(json_path)
        if not documents:
            print("❌ JSON verisi yüklenemedi veya boş")
            return
        
        split_docs = builder.split_documents(documents)
        vectorstore.add_documents(split_docs)
        
        # Kaydet
        vectorstore.save_local(vs_path)
        print("✅ Belgeler vectorstore'a eklendi ve kaydedildi.")


def main():
    parser = argparse.ArgumentParser(description="Belgeleri vectorstore'a ekleme aracı")
    parser.add_argument("--documents-dir", default="belgeler", 
                       help="Belgeler klasörü yolu (varsayılan: belgeler)")
    parser.add_argument("--base-url", default="https://alternativkraft.com/",
                       help="Base URL (varsayılan: https://alternativkraft.com/)")
    parser.add_argument("--json-only", action="store_true",
                       help="Sadece JSON oluştur, vectorstore'a ekleme")
    parser.add_argument("--rebuild", action="store_true",
                       help="Vectorstore'u baştan oluştur")
    parser.add_argument("--json-file", default=None,
                       help="Belirli bir JSON dosyasını kullan")
    
    args = parser.parse_args()
    
    # Belgeleri işle
    if args.json_file and os.path.exists(args.json_file):
        print(f"📄 Mevcut JSON kullanılıyor: {args.json_file}")
        json_path = args.json_file
    else:
        documents = process_documents(args.documents_dir, args.base_url)
        if not documents:
            print("❌ İşlenecek belge bulunamadı!")
            return
        
        # JSON'a kaydet
        json_path = save_documents_json(documents)
    
    if args.json_only:
        print("⏹️  --json-only seçildi. Sadece JSON oluşturuldu.")
        return
    
    # Vectorstore'a ekle
    add_to_vectorstore(json_path, rebuild=args.rebuild)


if __name__ == "__main__":
    main()
