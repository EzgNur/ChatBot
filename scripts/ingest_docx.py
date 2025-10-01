import sys
import pathlib
import requests
from typing import Optional
import argparse

try:
    from docx import Document  # python-docx
except Exception as e:
    raise SystemExit("python-docx paketi gerekli. Kurulum: ./chatbot_env/bin/pip install python-docx")


def docx_to_text(path: pathlib.Path) -> str:
    doc = Document(str(path))
    parts = []
    # Paragraphs (başlıkları ve listeleri koru)
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            # Vectorstore chunking separatorlarıyla uyumlu başlıklar
            level = 2
            try:
                level = int(''.join(ch for ch in p.style.name if ch.isdigit()) or '2')
            except Exception:
                level = 2
            parts.append(("#" * max(2, min(6, level))) + f" {txt}")
        elif style.startswith("list") or txt.startswith(("•", "-", "*")):
            # Madde işaretlerini normalize et
            clean = txt.lstrip("•*- ").strip()
            parts.append(f"- {clean}")
        else:
            parts.append(txt)
    # Tables (varsa)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def infer_title(file_path: pathlib.Path) -> str:
    title = file_path.stem.replace("_", " ").replace("-", " ").strip()
    # İlk harfleri büyüt
    return " ".join(w.capitalize() for w in title.split()) or "Belge"


def ingest_text(
    base_url: str,
    text: str,
    title: str,
    url: Optional[str] = "",
    author: Optional[str] = "Şirket",
    clean: bool = True,
) -> requests.Response:
    data = {
        # Başlığı içeriğin başına da ekleyerek aramada sinyali artır
        "text": (f"Başlık: {title}\n\n" + text).strip(),
        "title": title,
        "url": url or "",
        "author": author or "Şirket",
        "clean": str(bool(clean)).lower(),
    }
    return requests.post(f"{base_url}/ingest/transcript", data=data, timeout=120)


def main() -> None:
    parser = argparse.ArgumentParser(description="DOCX ingest aracı")
    parser.add_argument("folder", help="DOCX klasörü (mutlak yol)")
    parser.add_argument("api_base", nargs="?", default="http://127.0.0.1:8000", help="API base URL (vars: http://127.0.0.1:8000)")
    parser.add_argument("url_prefix", nargs="?", default="", help="URL prefix veya '@' ile sabit URL")
    parser.add_argument("--only", dest="only", default=None, help="Sadece bu dosya adını işle (tam dosya adı, .docx dahil)")
    parser.add_argument("--url", dest="fixed_url", default=None, help="Yalnızca --only ile birlikte: bu sabit URL'yi kullan")
    args = parser.parse_args()

    folder = pathlib.Path(args.folder).expanduser()
    base_url = args.api_base
    url_prefix = args.url_prefix

    if not folder.exists() or not folder.is_dir():
        print(f"Klasör bulunamadı: {folder}")
        sys.exit(1)

    if args.only:
        target = folder / args.only
        docx_files = [target] if target.exists() else []
    else:
        docx_files = sorted(folder.glob("*.docx"))
    if not docx_files:
        print("Hiç .docx dosyası bulunamadı.")
        sys.exit(0)

    print(f"Bulunan .docx dosyaları: {len(docx_files)}")

    ok, fail = 0, 0
    import unicodedata, re

    def slugify(name: str) -> str:
        s = unicodedata.normalize('NFKD', name)
        s = ''.join(ch for ch in s if not unicodedata.category(ch).startswith('M'))
        s = s.lower().strip().replace(' ', '-')
        s = re.sub(r"[^a-z0-9\-]", "", s)
        s = re.sub(r"-+", "-", s).strip('-')
        return s or "belge"

    for fp in docx_files:
        try:
            title = infer_title(fp)
            text = docx_to_text(fp)
            if not text:
                print(f"[SKIP] Boş içerik: {fp.name}")
                continue
            # URL üretim:
            # - Eğer url_prefix '@' ile başlıyorsa sabit URL kullan (slug ekleme)
            # - Aksi halde prefix + slug uygula
            gen_url = ""
            if args.fixed_url:
                gen_url = args.fixed_url
            elif url_prefix:
                if url_prefix.startswith('@'):
                    gen_url = url_prefix[1:]
                else:
                    gen_url = url_prefix.rstrip('/') + "/" + slugify(fp.stem)
            r = ingest_text(base_url, text, title, url=gen_url, author="Şirket", clean=True)
            if r.ok:
                ok += 1
                print(f"[OK] {fp.name} → {title}")
            else:
                fail += 1
                print(f"[FAIL] {fp.name} → HTTP {r.status_code} {r.text[:200]}")
        except Exception as e:
            fail += 1
            print(f"[ERR] {fp.name} → {e}")

    print(f"\nÖZET: başarı={ok} hata={fail}")


if __name__ == "__main__":
    main()


